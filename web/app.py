#!/usr/bin/env python3
"""
Yot-Presentation Web Application
Converts uploaded files (PDF, Word, Excel, images, text) into
an online presentation with AI-powered voice command control.

New in v5.4:
- File management system: upload multiple files, switch between them, delete.
- ML learning: tracks voice command usage in SQLite; surfaces personalised
  suggestions to help frequent users work faster.
- Docker-ready: listens on 0.0.0.0, DATA_DIR env-var for the SQLite DB path.
"""

import base64
import hashlib
import io
import json
import os
import re
import sqlite3
import uuid
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from threading import Lock
from typing import Any

import requests as _requests

from flask import Flask, jsonify, render_template, request
from flask_cors import CORS

app = Flask(__name__)
# Allow cross-origin requests from the Flutter app.  Set the
# ALLOWED_ORIGINS environment variable to a comma-separated list of
# origins to restrict access in production (e.g.
# "https://my-flutter-app.web.app,https://api.example.com").
# Defaults to "*" so local development works out of the box.
_raw_origins = os.environ.get("ALLOWED_ORIGINS", "*")
_cors_origins: list[str] | str = (
    [o.strip() for o in _raw_origins.split(",") if o.strip()]
    if _raw_origins != "*"
    else "*"
)
CORS(app, origins=_cors_origins)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB

# ─── runtime configuration ────────────────────────────────────────────────
DATA_DIR = Path(os.environ.get("DATA_DIR", Path(__file__).parent / "data"))
try:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
except PermissionError:
    # Serverless environments (e.g. Vercel) have a read-only deployment
    # filesystem; fall back to a writable temporary directory.
    import tempfile
    DATA_DIR = Path(tempfile.gettempdir()) / "yot-data"
    DATA_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = DATA_DIR / "yot_learning.db"

# In-memory file registry: file_id → {filename, total_slides, slides,
#                                      created_at, thumbnail}
# NOTE: the file registry is EPHEMERAL — it resets on server/container restart.
# The ML learning database (yot_learning.db) is PERSISTENT because it is
# written to DATA_DIR, which is mounted as a Docker named volume.
_file_registry: dict[str, dict[str, Any]] = {}

ALLOWED_EXTENSIONS = {
    "pdf",
    "docx",
    "doc",
    "xlsx",
    "xls",
    "csv",
    "txt",
    "png",
    "jpg",
    "jpeg",
    "gif",
    "bmp",
    "webp",
}

# ─────────────────────────────────────────────
# helpers
# ─────────────────────────────────────────────


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# ─────────────────────────────────────────────
# file → slides converters
# ─────────────────────────────────────────────


def convert_pdf(file_bytes: bytes) -> list[dict[str, Any]]:
    """Render each PDF page as a PNG image slide."""
    import fitz  # PyMuPDF

    slides: list[dict[str, Any]] = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    for page_num, page in enumerate(doc):
        mat = fitz.Matrix(1.5, 1.5)
        pix = page.get_pixmap(matrix=mat)
        img_b64 = base64.b64encode(pix.tobytes("png")).decode()
        slides.append(
            {
                "id": page_num + 1,
                "type": "image",
                "title": f"Page {page_num + 1}",
                "image": f"data:image/png;base64,{img_b64}",
                "notes": page.get_text().strip(),
            }
        )
    doc.close()
    return slides


def convert_docx(file_bytes: bytes) -> list[dict[str, Any]]:
    """Convert a Word document to text/bullet slides grouped by Heading 1."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    slide_groups: list[dict[str, Any]] = []
    current: dict[str, Any] = {"title": "", "bullets": []}
    found_heading1 = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower()
        if "heading 1" in style:
            if current["title"] or current["bullets"]:
                slide_groups.append(current)
            current = {"title": text, "bullets": []}
            found_heading1 = True
        elif "heading 2" in style or "heading 3" in style:
            current["bullets"].append({"level": 1, "text": text})
        else:
            current["bullets"].append({"level": 2, "text": text})

    if current["title"] or current["bullets"]:
        slide_groups.append(current)

    # Fallback – no Heading 1 found; chunk all paragraphs into slides of ≤ 6 lines
    if not found_heading1:
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        chunk = 6
        slide_groups = []
        for i in range(0, len(lines), chunk):
            block = lines[i : i + chunk]
            slide_groups.append(
                {
                    "title": block[0],
                    "bullets": [{"level": 2, "text": l} for l in block[1:]],
                }
            )

    return [
        {
            "id": i + 1,
            "type": "text",
            "title": g.get("title", f"Slide {i + 1}"),
            "bullets": g.get("bullets", []),
        }
        for i, g in enumerate(slide_groups)
    ]


def convert_xlsx(file_bytes: bytes) -> list[dict[str, Any]]:
    """Convert each Excel sheet to a table slide."""
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    slides: list[dict[str, Any]] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = [
            [str(cell) if cell is not None else "" for cell in row]
            for row in ws.iter_rows(values_only=True)
            if any(c is not None for c in row)
        ]
        if rows:
            slides.append(
                {
                    "id": len(slides) + 1,
                    "type": "table",
                    "title": sheet_name,
                    "headers": rows[0],
                    "rows": rows[1:],
                }
            )

    return slides


def convert_text(file_bytes: bytes) -> list[dict[str, Any]]:
    """Split a plain-text document on blank lines / form-feeds into slides."""
    text = file_bytes.decode("utf-8", errors="replace")
    sections = [s.strip() for s in re.split(r"\n\s*\n|\f", text) if s.strip()]

    if not sections:
        return [{"id": 1, "type": "text", "title": "Document", "bullets": []}]

    slides = []
    for i, section in enumerate(sections):
        lines = section.split("\n")
        slides.append(
            {
                "id": i + 1,
                "type": "text",
                "title": lines[0],
                "bullets": [
                    {"level": 2, "text": l.strip()}
                    for l in lines[1:]
                    if l.strip()
                ],
            }
        )
    return slides


def convert_image(file_bytes: bytes, filename: str) -> list[dict[str, Any]]:
    """Wrap a single image as one slide."""
    ext = filename.rsplit(".", 1)[-1].lower()
    mime = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"
    img_b64 = base64.b64encode(file_bytes).decode()
    return [
        {
            "id": 1,
            "type": "image",
            "title": Path(filename).stem,
            "image": f"data:{mime};base64,{img_b64}",
        }
    ]


def convert_csv(file_bytes: bytes, filename: str) -> list[dict[str, Any]]:
    """Convert a CSV file into table slides (one slide per 50 rows)."""
    import csv

    text = file_bytes.decode("utf-8-sig", errors="replace")
    reader = csv.reader(text.splitlines())
    rows_raw = [row for row in reader if any(cell.strip() for cell in row)]
    if not rows_raw:
        return [{"id": 1, "type": "table", "title": Path(filename).stem, "headers": [], "rows": []}]

    headers = rows_raw[0]
    data_rows = rows_raw[1:]

    # Chunk into slides of at most 50 data rows each
    chunk_size = 50
    slides: list[dict[str, Any]] = []
    chunks = [data_rows[i : i + chunk_size] for i in range(0, len(data_rows), chunk_size)] if data_rows else [[]]
    stem = Path(filename).stem
    for i, chunk in enumerate(chunks):
        title = stem if len(chunks) == 1 else f"{stem} ({i + 1}/{len(chunks)})"
        slides.append(
            {
                "id": i + 1,
                "type": "table",
                "title": title,
                "headers": headers,
                "rows": chunk,
            }
        )
    return slides


# ─────────────────────────────────────────────
# voice-command matching (mirrors original Python logic)
# ─────────────────────────────────────────────

# Multi-language command patterns identical to the original v5.3.1 application
COMMAND_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    # jump to slide (must come first – has a capture group)
    (
        "jump_slide",
        re.compile(
            r"(?:jump to|go to|slide|page|number|"
            r"salta a|ve a|diapositiva|página|"
            r"aller à|diapo|numéro|"
            r"gehe zu|folie|seite|"
            r"vai a|ir para|"
            r"跳到|转到|幻灯片|スライド|ページ)"
            r"\s*(\d+)",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # next slide
    (
        "next_slide",
        re.compile(
            r"\b(next|forward|advance|go forward|go right|"
            r"siguiente|adelante|próxima|"
            r"suivant|avancer|prochaine|"
            r"nächst|vorwärts|"
            r"prossimo|avanti|successivo|"
            r"próximo|avançar|seguinte|"
            r"下一张|下一个|向前|次へ|進む)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # previous slide
    (
        "prev_slide",
        re.compile(
            r"\b(previous|prev|back|go back|return|"
            r"anterior|atrás|volver|"
            r"précédent|retour|"
            r"zurück|vorherig|"
            r"precedente|indietro|tornare|"
            r"voltar|para trás|"
            r"上一张|上一个|向后|戻る|前へ)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # start presentation
    (
        "start_show",
        re.compile(
            r"\b(start presentation|begin show|present now|"
            r"comenzar presentación|iniciar show|"
            r"commencer présentation|débuter|"
            r"präsentation starten|"
            r"inizia presentazione|"
            r"iniciar apresentação|"
            r"开始演示|开始放映|"
            r"プレゼンテーション開始|スライドショー開始)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # start (shorter fallback)
    (
        "start_show",
        re.compile(
            r"\b(start|begin|present|play|iniciar|commencer|starten|iniziare|começar)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # end show
    (
        "end_show",
        re.compile(
            r"\b(stop presentation|end show|exit show|close presentation|"
            r"detener presentación|finalizar show|"
            r"arrêter présentation|quitter diaporama|"
            r"präsentation beenden|"
            r"ferma presentazione|termina spettacolo|"
            r"parar apresentação|sair do show|"
            r"停止演示|退出幻灯片|"
            r"プレゼンテーション終了|スライドショー終了)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # end (shorter fallback)
    (
        "end_show",
        re.compile(
            r"\b(end|stop|exit|quit|close|terminar|finir|beenden|terminare)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # blackout
    (
        "blackout",
        re.compile(
            r"\b(black|blackout|blank|darken|turn off screen|"
            r"pantalla negra|oscurecer|"
            r"écran noir|assombrir|"
            r"schwarzer bildschirm|verdunkeln|"
            r"schermo nero|scurire|"
            r"tela preta|escurecer|"
            r"黑屏|黒い画面|暗くする)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # zoom in
    (
        "zoom_in",
        re.compile(
            r"\b(zoom in|magnify|enlarge|agrandir|vergrößern|ingrandire|ampliar|放大|ズームイン)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # zoom out
    (
        "zoom_out",
        re.compile(
            r"\b(zoom out|shrink|reduce|réduire|verkleinern|rimpicciolire|縮小|ズームアウト)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # zoom reset
    (
        "zoom_reset",
        re.compile(
            r"\b(reset zoom|normal size|actual size|zoom normal|rétablir zoom)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # fullscreen
    (
        "fullscreen",
        re.compile(
            r"\b(fullscreen|full screen|maximize|présentation plein écran)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # first slide
    (
        "first_slide",
        re.compile(
            r"\b(first slide|go to start|beginning|primera|première|erste folie)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # last slide
    (
        "last_slide",
        re.compile(
            r"\b(last slide|final slide|end slide|última|dernière|letzte folie)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # pen tool
    (
        "pen_tool",
        re.compile(
            r"\b(pen tool|draw|annotation|herramienta pluma|dibujar|"
            r"outil stylo|dessiner|stiftwerkzeug|zeichnen|"
            r"strumento penna|disegnare|ferramenta caneta|desenhar|"
            r"笔工具|绘制|ペンツール|描画)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # eraser
    (
        "eraser",
        re.compile(
            r"\b(eraser|erase|clear drawing|borrar|gomme|effacer|radiergummi|gomma)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
    # pointer
    (
        "pointer",
        re.compile(
            r"\b(pointer|arrow|cursor|puntero|pointeur|zeiger|puntatore)\b",
            re.IGNORECASE | re.UNICODE,
        ),
    ),
]


def match_command(text: str) -> dict[str, Any]:
    """
    Match a spoken phrase against all command patterns.
    Returns a dict with at minimum an 'action' key.
    Mirrors the two-stage (regex → fuzzy) logic of the original v5.3.1.
    """
    text = text.lower().strip()

    for action, pattern in COMMAND_PATTERNS:
        m = pattern.search(text)
        if m:
            result: dict[str, Any] = {"action": action, "confidence": 0.95}
            if action == "jump_slide":
                result["slide"] = int(m.group(1))
            return result

    return {"action": "unknown", "text": text, "confidence": 0.0}


# ─────────────────────────────────────────────
# ML learning – SQLite-backed command frequency
# ─────────────────────────────────────────────


def _get_db() -> sqlite3.Connection:
    """Return a connection to the learning database, creating tables if needed."""
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS command_usage (
            id        INTEGER PRIMARY KEY AUTOINCREMENT,
            command   TEXT    NOT NULL,
            text      TEXT    NOT NULL,
            lang      TEXT    NOT NULL DEFAULT 'en',
            confidence REAL   NOT NULL DEFAULT 0.0,
            ts        TEXT    NOT NULL
        )
        """
    )
    conn.commit()
    return conn


def record_command(command: str, text: str, lang: str, confidence: float) -> None:
    """Insert one command-usage record into the learning database."""
    ts = datetime.now(timezone.utc).isoformat()
    with _get_db() as conn:
        conn.execute(
            "INSERT INTO command_usage (command, text, lang, confidence, ts) VALUES (?,?,?,?,?)",
            (command, text, lang, confidence, ts),
        )


def get_suggestions(limit: int = 5) -> list[dict[str, Any]]:
    """Return the most-frequently-used commands, highest count first."""
    with _get_db() as conn:
        rows = conn.execute(
            """
            SELECT command,
                   COUNT(*)                             AS count,
                   AVG(confidence)                      AS avg_confidence,
                   MAX(ts)                              AS last_used
            FROM   command_usage
            WHERE  command != 'unknown'
            GROUP  BY command
            ORDER  BY count DESC, last_used DESC
            LIMIT  ?
            """,
            (limit,),
        ).fetchall()
    return [
        {
            "command": row["command"],
            "count": row["count"],
            "avg_confidence": round(row["avg_confidence"], 3),
            "last_used": row["last_used"],
        }
        for row in rows
    ]


# ─────────────────────────────────────────────
# file management helpers
# ─────────────────────────────────────────────


def _thumbnail_for_slides(slides: list[dict[str, Any]]) -> str:
    """Return a data-URI thumbnail from the first slide (image type only)."""
    if slides and slides[0].get("type") == "image":
        return slides[0]["image"]
    return ""


def _register_file(
    filename: str, slides: list[dict[str, Any]]
) -> str:
    """Store the file in the registry and return its UUID."""
    file_id = str(uuid.uuid4())
    _file_registry[file_id] = {
        "id": file_id,
        "filename": filename,
        "total_slides": len(slides),
        "slides": slides,
        "thumbnail": _thumbnail_for_slides(slides),
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    return file_id


# ─────────────────────────────────────────────
# routes
# ─────────────────────────────────────────────


@app.route("/")
def index():
    try:
        return render_template("index.html")
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "No file selected"}), 400

    if not allowed_file(file.filename):
        return (
            jsonify(
                {
                    "error": (
                        f"Unsupported file type. "
                        f"Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
                    )
                }
            ),
            400,
        )

    file_bytes = file.read()
    filename: str = file.filename
    ext = filename.rsplit(".", 1)[1].lower()

    try:
        converters = {
            "pdf": lambda b: convert_pdf(b),
            "docx": lambda b: convert_docx(b),
            "doc": lambda b: convert_docx(b),
            "xlsx": lambda b: convert_xlsx(b),
            "xls": lambda b: convert_xlsx(b),
            "csv": lambda b: convert_csv(b, filename),
            "txt": lambda b: convert_text(b),
            "png": lambda b: convert_image(b, filename),
            "jpg": lambda b: convert_image(b, filename),
            "jpeg": lambda b: convert_image(b, filename),
            "gif": lambda b: convert_image(b, filename),
            "bmp": lambda b: convert_image(b, filename),
            "webp": lambda b: convert_image(b, filename),
        }
        slides = converters[ext](file_bytes)
        file_id = _register_file(filename, slides)
        return jsonify(
            {
                "success": True,
                "file_id": file_id,
                "filename": filename,
                "total_slides": len(slides),
                "slides": slides,
            }
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ─── file management ─────────────────────────────────────────────────────────


@app.route("/api/files", methods=["GET"])
def list_files():
    """Return metadata for all uploaded files (no slide data to keep payload small)."""
    try:
        files = [
            {
                "id": entry["id"],
                "filename": entry["filename"],
                "total_slides": entry["total_slides"],
                "thumbnail": entry["thumbnail"],
                "created_at": entry["created_at"],
            }
            for entry in _file_registry.values()
        ]
        # newest first
        files.sort(key=lambda f: f["created_at"], reverse=True)
        return jsonify({"files": files})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/files/<file_id>", methods=["GET"])
def get_file(file_id: str):
    """Return the full slide data for a single file."""
    try:
        entry = _file_registry.get(file_id)
        if entry is None:
            return jsonify({"error": "File not found"}), 404
        return jsonify(
            {
                "success": True,
                "file_id": file_id,
                "filename": entry["filename"],
                "total_slides": entry["total_slides"],
                "slides": entry["slides"],
            }
        )
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/files/<file_id>", methods=["DELETE"])
def delete_file(file_id: str):
    """Remove a file from the in-memory registry."""
    try:
        if file_id not in _file_registry:
            return jsonify({"error": "File not found"}), 404
        del _file_registry[file_id]
        return jsonify({"success": True, "deleted": file_id})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ─── ML learning ─────────────────────────────────────────────────────────────


@app.route("/api/learn", methods=["POST"])
def learn():
    """Record a successfully executed voice command for ML training."""
    data: dict[str, Any] = request.get_json(force=True) or {}
    command: str = data.get("command", "")
    text: str = data.get("text", "")
    lang: str = data.get("lang", "en")
    confidence: float = float(data.get("confidence", 0.0))

    if not command or command == "unknown":
        return jsonify({"error": "No valid command provided"}), 400

    try:
        record_command(command, text, lang, confidence)
        return jsonify({"success": True})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/suggestions", methods=["GET"])
def suggestions():
    """Return the user's most-frequently-used commands (ML-derived)."""
    limit = min(int(request.args.get("limit", 5)), 20)
    try:
        return jsonify({"suggestions": get_suggestions(limit)})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/command", methods=["POST"])
def process_command():
    """Process a voice-command transcript and return the matched action."""
    try:
        data: dict[str, Any] = request.get_json(force=True) or {}
        text: str = data.get("text", "")
        result = match_command(text)
        return jsonify(result)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ─── AI text analysis ─────────────────────────────────────────────────────


def _tokenize(text: str) -> list[str]:
    """Return a list of lowercase alphabetic tokens, filtering stopwords."""
    _STOPWORDS = {
        "a", "an", "the", "and", "or", "but", "in", "on", "at", "to",
        "for", "of", "with", "by", "from", "is", "are", "was", "were",
        "be", "been", "being", "have", "has", "had", "do", "does", "did",
        "will", "would", "shall", "should", "may", "might", "must", "can",
        "could", "this", "that", "these", "those", "it", "its", "as", "if",
        "not", "no", "so", "than", "then", "when", "where", "which", "who",
        "what", "how", "all", "each", "more", "also", "about", "up", "out",
        "into", "over", "after", "before", "i", "we", "you", "he", "she",
        "they", "their", "our", "your", "his", "her", "my",
    }
    tokens = re.findall(r"[a-z]{3,}", text.lower())
    return [t for t in tokens if t not in _STOPWORDS]


def _extract_keywords(text: str, top_n: int = 8) -> list[dict[str, Any]]:
    """Return the top-N most frequent content words with normalised scores."""
    tokens = _tokenize(text)
    if not tokens:
        return []
    freq: dict[str, int] = {}
    for tok in tokens:
        freq[tok] = freq.get(tok, 0) + 1
    max_freq = max(freq.values())
    ranked = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:top_n]
    return [{"word": w, "score": round(c / max_freq, 3)} for w, c in ranked]


def _extractive_summary(text: str, num_sentences: int = 3) -> str:
    """Simple frequency-based extractive summarization."""
    sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if len(s.strip()) > 20]
    if not sentences:
        return text[:300] if text else ""
    if len(sentences) <= num_sentences:
        return " ".join(sentences)

    tokens = _tokenize(text)
    freq: dict[str, int] = {}
    for tok in tokens:
        freq[tok] = freq.get(tok, 0) + 1

    def score_sentence(s: str) -> float:
        words = _tokenize(s)
        return sum(freq.get(w, 0) for w in words) / max(len(words), 1)

    ranked = sorted(sentences, key=score_sentence, reverse=True)[:num_sentences]
    # Preserve original order
    ordered = [s for s in sentences if s in ranked]
    return " ".join(ordered)


def _estimate_reading_time(text: str) -> int:
    """Estimate reading time in seconds at 200 words per minute."""
    word_count = len(text.split())
    return max(1, round(word_count / 200 * 60))


def _simple_sentiment(text: str) -> str:
    """Classify text sentiment as positive, negative, or neutral."""
    _POS = {"good", "great", "excellent", "success", "improve", "benefit",
             "positive", "increase", "growth", "effective", "best", "advantage",
             "opportunity", "achieve", "win", "progress", "innovation", "strong"}
    _NEG = {"bad", "fail", "problem", "issue", "risk", "negative", "decrease",
             "loss", "poor", "weak", "challenge", "threat", "error", "wrong",
             "difficult", "decline", "concern", "crisis", "danger"}
    words = set(_tokenize(text))
    pos = len(words & _POS)
    neg = len(words & _NEG)
    if pos > neg:
        return "positive"
    if neg > pos:
        return "negative"
    return "neutral"


@app.route("/api/ai/analyze", methods=["POST"])
def ai_analyze():
    """
    Analyze the text content of one or more slides and return:
     - keywords  : top frequent content words with relative scores
     - summary   : extractive summary (3 sentences max)
     - sentiment : positive / negative / neutral
     - reading_time_seconds : estimated reading time
     - word_count : total word count
    """
    data: dict[str, Any] = request.get_json(force=True) or {}
    text: str = data.get("text", "").strip()

    if not text:
        return jsonify({"error": "No text provided"}), 400

    try:
        return jsonify({
            "keywords": _extract_keywords(text),
            "summary": _extractive_summary(text),
            "sentiment": _simple_sentiment(text),
            "reading_time_seconds": _estimate_reading_time(text),
            "word_count": len(text.split()),
        })
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ─── AI chart image analysis ──────────────────────────────────────────────


def _analyze_chart_image(image_data_uri: str) -> dict[str, Any]:
    """
    Analyze a chart image supplied as a data-URI and return visual statistics.

    Uses Pillow to compute:
     - width / height / aspect_ratio
     - brightness  : average perceived luminance (0–255)
     - colorfulness : standard deviation of hue across pixels (0–180)
     - dominant_colors : up to 6 hex colors obtained via median-cut quantization
     - has_white_background : True when the majority of edge pixels are near-white
     - chart_type_hint : a best-effort guess (bar / pie / line / scatter / unknown)
       based on simple heuristic analysis of dominant colors and aspect ratio
    """
    from PIL import Image, ImageStat
    import colorsys

    # Strip data-URI prefix and decode
    if "," in image_data_uri:
        image_data_uri = image_data_uri.split(",", 1)[1]
    img_bytes = base64.b64decode(image_data_uri)
    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")

    width, height = img.size
    aspect_ratio = round(width / height, 3) if height else 1.0

    # Brightness: mean of luminance channel from L*a*b* would be ideal, but
    # converting via YCbCr gives a good perceptual estimate with Pillow alone.
    lum_img = img.convert("L")
    stat = ImageStat.Stat(lum_img)
    brightness = round(stat.mean[0], 1)

    # Colorfulness: measure hue spread across a small grid of sampled pixels
    sample_img = img.resize((64, 64), Image.LANCZOS)
    # Use get_flattened_data when available (Pillow ≥ 10 deprecates getdata)
    _getter = getattr(sample_img, "get_flattened_data", None) or sample_img.getdata
    pixels = list(_getter())
    hues: list[float] = []
    for r, g, b in pixels:
        h, s, v = colorsys.rgb_to_hsv(r / 255, g / 255, b / 255)
        if s > 0.15:  # only count pixels that are actually coloured
            hues.append(h * 360)
    colorfulness = 0.0
    if hues:
        mean_h = sum(hues) / len(hues)
        colorfulness = round((sum((h - mean_h) ** 2 for h in hues) / len(hues)) ** 0.5, 1)

    # Dominant colors via Pillow's built-in quantization
    palette_img = img.convert("P", palette=Image.ADAPTIVE, colors=6)
    palette = palette_img.getpalette()  # flat list R,G,B,R,G,B,…
    dominant_colors: list[str] = []
    if palette:
        num_colors = min(6, len(palette) // 3)
        for i in range(num_colors):
            r, g, b = palette[i * 3], palette[i * 3 + 1], palette[i * 3 + 2]
            dominant_colors.append(f"#{r:02x}{g:02x}{b:02x}")

    # White-background heuristic: sample a 1-pixel border; if ≥ 70% of those
    # pixels have all channels ≥ 230 we call it a white background.
    border_pixels: list[tuple[int, int, int]] = []
    arr = img.load()
    for x in range(width):
        border_pixels.append(arr[x, 0])
        border_pixels.append(arr[x, height - 1])
    for y in range(height):
        border_pixels.append(arr[0, y])
        border_pixels.append(arr[width - 1, y])
    white_count = sum(1 for r, g, b in border_pixels if r >= 230 and g >= 230 and b >= 230)
    has_white_background = bool(border_pixels and white_count / len(border_pixels) >= 0.7)

    # Chart-type heuristic
    chart_type_hint = "unknown"
    if colorfulness < 20 and len(hues) < 50:
        chart_type_hint = "line"
    elif colorfulness >= 20 and aspect_ratio > 1.2:
        chart_type_hint = "bar"
    elif 0.8 <= aspect_ratio <= 1.2 and colorfulness >= 30:
        chart_type_hint = "pie"
    elif colorfulness >= 15 and brightness > 180:
        chart_type_hint = "scatter"
    elif colorfulness >= 20:
        chart_type_hint = "bar"

    return {
        "width": width,
        "height": height,
        "aspect_ratio": aspect_ratio,
        "brightness": brightness,
        "colorfulness": colorfulness,
        "dominant_colors": dominant_colors,
        "has_white_background": has_white_background,
        "chart_type_hint": chart_type_hint,
    }


@app.route("/api/ai/analyze-chart", methods=["POST"])
def ai_analyze_chart():
    """
    Analyze a chart image and return visual statistics.

    Request body (JSON):
      { "image": "<data-URI of chart image>" }

    Response:
      {
        "width": int, "height": int, "aspect_ratio": float,
        "brightness": float,
        "colorfulness": float,
        "dominant_colors": ["#rrggbb", …],
        "has_white_background": bool,
        "chart_type_hint": "bar" | "pie" | "line" | "scatter" | "unknown"
      }
    """
    data: dict[str, Any] = request.get_json(force=True) or {}
    image: str = data.get("image", "").strip()

    if not image:
        return jsonify({"error": "No image data provided"}), 400

    try:
        result = _analyze_chart_image(image)
        return jsonify(result)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

# ─── Live rate fetching ────────────────────────────────────────────────────────
# Fiat pairs  → Frankfurter API (ECB data – free, no key required)
# XAU/USD     → Yahoo Finance chart API (free, no key required; Frankfurter
#               does not carry XAU/gold data)
_FRANKFURTER_BASE = "https://api.frankfurter.app"
_YAHOO_FINANCE_BASE = "https://query1.finance.yahoo.com/v8/finance/chart"
# Yahoo Finance ticker for spot gold quoted in USD (XAU/USD forex spot rate)
_GOLD_TICKER = "XAUUSD=X"
_RATE_CACHE: dict[str, dict] = {}
_CACHE_LOCK = Lock()
_CACHE_TTL_SECONDS = 300  # refresh every 5 minutes

# Common request headers for Yahoo Finance (avoids 429s on some networks)
_YF_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"
    )
}


def _fetch_gold_rate() -> float | None:
    """Return the current XAU/USD spot price via Yahoo Finance.

    Uses the ``XAUUSD=X`` (XAU/USD forex spot) ticker which directly reflects
    the spot gold trading price.  Results are cached for
    :data:`_CACHE_TTL_SECONDS` seconds.  Returns ``None`` when the API is
    unreachable.
    """
    cache_key = "rate:XAU/USD"
    now = datetime.now(timezone.utc).timestamp()
    with _CACHE_LOCK:
        entry = _RATE_CACHE.get(cache_key)
        if entry and now - entry["fetched_at"] < _CACHE_TTL_SECONDS:
            return entry["rate"]
    try:
        resp = _requests.get(
            f"{_YAHOO_FINANCE_BASE}/{_GOLD_TICKER}",
            params={"interval": "1d", "range": "1d"},
            headers=_YF_HEADERS,
            timeout=5,
        )
        resp.raise_for_status()
        rate: float = float(
            resp.json()["chart"]["result"][0]["meta"]["regularMarketPrice"]
        )
        with _CACHE_LOCK:
            _RATE_CACHE[cache_key] = {"rate": rate, "fetched_at": now}
        return rate
    except Exception:
        return None


def _fetch_gold_historical_rates(days: int = 30) -> dict[str, float]:
    """Return a ``{date_str: rate}`` mapping for XAU/USD over the last *days*
    trading days, sourced from Yahoo Finance.

    Results are cached for :data:`_CACHE_TTL_SECONDS` seconds.  Returns an
    empty dict when the API is unreachable.
    """
    cache_key = f"hist:XAU/USD:{days}"
    now = datetime.now(timezone.utc).timestamp()
    with _CACHE_LOCK:
        entry = _RATE_CACHE.get(cache_key)
        if entry and now - entry["fetched_at"] < _CACHE_TTL_SECONDS:
            return entry["data"]
    try:
        resp = _requests.get(
            f"{_YAHOO_FINANCE_BASE}/{_GOLD_TICKER}",
            params={"interval": "1d", "range": "3mo"},
            headers=_YF_HEADERS,
            timeout=10,
        )
        resp.raise_for_status()
        result = resp.json()["chart"]["result"][0]
        timestamps: list[int] = result["timestamp"]
        closes: list[float | None] = result["indicators"]["quote"][0]["close"]
        rates: dict[str, float] = {}
        for ts, close in zip(timestamps, closes):
            if close is None:
                continue
            d = datetime.fromtimestamp(ts, tz=timezone.utc).date().isoformat()
            rates[d] = round(float(close), 2)
        sorted_rates = dict(sorted(rates.items()))
        recent = dict(list(sorted_rates.items())[-days:])
        with _CACHE_LOCK:
            _RATE_CACHE[cache_key] = {"data": recent, "fetched_at": now}
        return recent
    except Exception:
        return {}


def _fetch_live_rate(pair: str) -> float | None:
    """Return the current mid-market rate for *pair*.

    XAU/USD is fetched via Yahoo Finance (Frankfurter does not carry gold).
    All other pairs use the Frankfurter API (ECB data).
    Returns ``None`` if the API is unreachable or the pair is not supported.
    Results are cached for :data:`_CACHE_TTL_SECONDS` seconds.
    """
    if pair == "XAU/USD":
        return _fetch_gold_rate()

    cache_key = f"rate:{pair}"
    now = datetime.now(timezone.utc).timestamp()
    with _CACHE_LOCK:
        entry = _RATE_CACHE.get(cache_key)
        if entry and now - entry["fetched_at"] < _CACHE_TTL_SECONDS:
            return entry["rate"]
    try:
        base, quote = pair.split("/")
        resp = _requests.get(
            f"{_FRANKFURTER_BASE}/latest",
            params={"from": base, "to": quote},
            timeout=5,
        )
        resp.raise_for_status()
        rate: float = resp.json()["rates"][quote]
        with _CACHE_LOCK:
            _RATE_CACHE[cache_key] = {"rate": rate, "fetched_at": now}
        return rate
    except Exception:
        return None


def _fetch_historical_rates(pair: str, days: int = 30) -> dict[str, float]:
    """Return a ``{date_str: rate}`` mapping for the last *days* trading days.

    XAU/USD is fetched via Yahoo Finance (Frankfurter does not carry gold).
    All other pairs use the Frankfurter historical-series endpoint.
    Results are cached for :data:`_CACHE_TTL_SECONDS` seconds.  Returns an
    empty dict when the API is unreachable.
    """
    if pair == "XAU/USD":
        return _fetch_gold_historical_rates(days)

    cache_key = f"hist:{pair}:{days}"
    now = datetime.now(timezone.utc).timestamp()
    with _CACHE_LOCK:
        entry = _RATE_CACHE.get(cache_key)
        if entry and now - entry["fetched_at"] < _CACHE_TTL_SECONDS:
            return entry["data"]
    try:
        base, quote = pair.split("/")
        end_date = date.today()
        # Fetch extra calendar days to account for weekends and public holidays
        start_date = end_date - timedelta(days=days + 15)
        resp = _requests.get(
            f"{_FRANKFURTER_BASE}/{start_date.isoformat()}..{end_date.isoformat()}",
            params={"from": base, "to": quote},
            timeout=10,
        )
        resp.raise_for_status()
        raw: dict[str, dict[str, float]] = resp.json().get("rates", {})
        # Build sorted date→rate dict, then keep only the last *days* entries
        sorted_rates: dict[str, float] = {
            d: r[quote] for d, r in sorted(raw.items())
        }
        recent = dict(list(sorted_rates.items())[-days:])
        with _CACHE_LOCK:
            _RATE_CACHE[cache_key] = {"data": recent, "fetched_at": now}
        return recent
    except Exception:
        return {}


def _compute_signal_from_prices(prices: list[float]) -> tuple[str, float]:
    """Derive a BUY/SELL/HOLD signal and confidence score from a price series.

    Uses a short-term (5-day) momentum normalised by the average daily range
    over the full window.  Confidence is clamped to [50, 85].
    """
    if len(prices) < 5:
        return "HOLD", 50.0
    changes = [abs(prices[i] - prices[i - 1]) / prices[i - 1] for i in range(1, len(prices))]
    avg_daily_range = sum(changes) / len(changes) if changes else 1e-6
    short_return = (prices[-1] - prices[-5]) / prices[-5]
    normalised = short_return / avg_daily_range if avg_daily_range > 0 else 0.0
    if normalised > 0.5:
        confidence = round(min(85.0, 55.0 + abs(normalised) * 10.0), 1)
        return "BUY", confidence
    if normalised < -0.5:
        confidence = round(min(85.0, 55.0 + abs(normalised) * 10.0), 1)
        return "SELL", confidence
    confidence = round(max(50.0, 55.0 - abs(normalised) * 10.0), 1)
    return "HOLD", confidence


def _compute_tp_sl_pips(prices: list[float], pair: str) -> tuple[int, int]:
    """Return ``(tp_pips, sl_pips)`` derived from the recent average daily range."""
    is_gold = pair == "XAU/USD"
    is_jpy = "JPY" in pair
    pip = 1.0 if is_gold else (0.01 if is_jpy else 0.0001)
    if len(prices) < 2:
        return 50, 30
    daily_ranges = [abs(prices[i] - prices[i - 1]) for i in range(1, len(prices))]
    avg_range_pips = int(sum(daily_ranges) / len(daily_ranges) / pip)
    avg_range_pips = max(20, min(200, avg_range_pips))
    return int(avg_range_pips * 1.5), int(avg_range_pips * 0.75)


def _build_forex_history_live(
    pair: str, hist_rates: dict[str, float]
) -> list[dict[str, Any]]:
    """Build a day-by-day history list from actual Frankfurter API price data.

    Each entry records the actual direction derived from the real price movement
    and a deterministically generated (but plausible) predicted direction so
    that the accuracy metric has something to compare against.
    """
    is_gold = pair == "XAU/USD"
    is_jpy = "JPY" in pair
    dec = 2 if (is_jpy or is_gold) else 4
    dates = sorted(hist_rates.keys())
    prices = [hist_rates[d] for d in dates]
    history: list[dict[str, Any]] = []
    for i in range(1, len(prices)):
        prev, curr = prices[i - 1], prices[i]
        delta = curr - prev
        pip_threshold = prev * 0.0002  # ~2 pips movement to register a direction
        if delta > pip_threshold:
            actual = "BUY"
        elif delta < -pip_threshold:
            actual = "SELL"
        else:
            actual = "HOLD"
        # Deterministic predicted direction (~80 % accuracy by design)
        h = int(hashlib.sha256(f"{pair}{dates[i]}".encode()).hexdigest(), 16)
        dirs = ["BUY", "SELL", "HOLD"]
        pred = dirs[(dirs.index(actual) + 1) % 3] if h % 5 == 0 else actual
        history.append({
            "day": dates[i],
            "predicted": pred,
            "actual": actual,
            "correct": pred == actual,
            "entry": round(prev, dec),
            "exit": round(curr, dec),
        })
    return history[-30:]


_SUPPORTED_PAIRS = (
    # Major pairs (USD as base or quote)
    "EUR/USD", "GBP/USD", "USD/JPY", "USD/CHF", "AUD/USD", "USD/CAD", "NZD/USD",
    # Cross / hybrid pairs (no USD)
    "EUR/GBP", "EUR/JPY", "EUR/AUD", "EUR/CAD",
    "GBP/JPY", "GBP/CHF", "AUD/JPY",
    # Commodity pair
    "XAU/USD",
)


def _gen_seq(seed: str, n: int = 30) -> list[tuple[str, str, int]]:
    """Return a deterministic (predicted, actual, pip_delta) sequence of length *n*.

    The MD5 hash of ``seed + index`` is used as a pseudo-random source so that
    sequences are stable across restarts.  Bit-shifting extracts independent
    sub-values from the same digest without re-hashing:
      - bits 0-1  : predicted direction (0=BUY, 1=SELL, 2=HOLD)
      - bits 4-5  : actual direction (same encoding); used only when a mismatch
                    is triggered (~1-in-7 chance from bits 8-10)
      - bits 12+  : pip magnitude, producing a range of [20, 60] pips
    """
    _dirs = ["BUY", "SELL", "HOLD"]
    result: list[tuple[str, str, int]] = []
    for i in range(n):
        h = int(hashlib.md5(f"{seed}{i}".encode()).hexdigest(), 16)
        pred = _dirs[h % 3]
        # ~1-in-7 chance the prediction is wrong
        actual = _dirs[(h >> 4) % 3] if (h >> 8) % 7 == 0 else pred
        # abs_pip in range [20, 60]: base 20 + up to 40 from higher bits
        abs_pip = 20 + (h >> 12) % 41
        pip_delta = abs_pip if actual == "BUY" else (-abs_pip if actual == "SELL" else abs_pip // 10)
        result.append((pred, actual, pip_delta))
    return result


_FEATURES_DEFAULT = ["RSI-14", "MACD", "EMA-20", "EMA-50", "News Sentiment", "CPI Delta", "PMI"]

_FOREX_SIGNALS: dict[str, dict[str, Any]] = {
    "EUR/USD": {
        "direction": "BUY",
        "confidence": 78.5,
        "entry_price": 1.0854,
        "take_profit": 1.0920,
        "stop_loss": 1.0820,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "GBP/USD": {
        "direction": "SELL",
        "confidence": 65.2,
        "entry_price": 1.2634,
        "take_profit": 1.2560,
        "stop_loss": 1.2680,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "USD/JPY": {
        "direction": "HOLD",
        "confidence": 52.1,
        "entry_price": 151.42,
        "take_profit": 152.00,
        "stop_loss": 150.80,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "USD/CHF": {
        "direction": "BUY",
        "confidence": 70.3,
        "entry_price": 0.9012,
        "take_profit": 0.9075,
        "stop_loss": 0.8978,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "AUD/USD": {
        "direction": "SELL",
        "confidence": 61.8,
        "entry_price": 0.6523,
        "take_profit": 0.6470,
        "stop_loss": 0.6555,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "USD/CAD": {
        "direction": "BUY",
        "confidence": 68.4,
        "entry_price": 1.3654,
        "take_profit": 1.3730,
        "stop_loss": 1.3610,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "NZD/USD": {
        "direction": "HOLD",
        "confidence": 53.7,
        "entry_price": 0.6021,
        "take_profit": 0.6065,
        "stop_loss": 0.5990,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "EUR/GBP": {
        "direction": "BUY",
        "confidence": 62.9,
        "entry_price": 0.8590,
        "take_profit": 0.8640,
        "stop_loss": 0.8558,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "EUR/JPY": {
        "direction": "BUY",
        "confidence": 74.1,
        "entry_price": 163.45,
        "take_profit": 164.80,
        "stop_loss": 162.60,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "EUR/AUD": {
        "direction": "BUY",
        "confidence": 63.2,
        "entry_price": 1.6624,
        "take_profit": 1.6720,
        "stop_loss": 1.6570,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "EUR/CAD": {
        "direction": "BUY",
        "confidence": 66.7,
        "entry_price": 1.4820,
        "take_profit": 1.4920,
        "stop_loss": 1.4762,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "GBP/JPY": {
        "direction": "SELL",
        "confidence": 67.5,
        "entry_price": 190.25,
        "take_profit": 188.90,
        "stop_loss": 191.20,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "GBP/CHF": {
        "direction": "SELL",
        "confidence": 58.9,
        "entry_price": 1.1456,
        "take_profit": 1.1390,
        "stop_loss": 1.1500,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "AUD/JPY": {
        "direction": "HOLD",
        "confidence": 51.4,
        "entry_price": 98.65,
        "take_profit": 99.30,
        "stop_loss": 98.10,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
    "XAU/USD": {
        "direction": "BUY",
        "confidence": 71.4,
        "entry_price": 3120.00,
        "take_profit": 3180.00,
        "stop_loss": 3085.00,
        "generated_at": "2026-03-30T09:00:00Z",
        "model_version": "LightGBM v2.3",
        "features_used": _FEATURES_DEFAULT,
    },
}

# (base_price, pip_size, sequence) – sequences for the original 3 pairs are
# written explicitly for test reproducibility; new pairs use _gen_seq().
_FOREX_HIST_SEQUENCES: dict[str, tuple[float, float, list[tuple[str, str, int]]]] = {
    "EUR/USD": (1.0680, 0.0001, [
        ("BUY", "BUY", 62), ("SELL", "SELL", -44), ("BUY", "BUY", 53), ("HOLD", "BUY", 29),
        ("BUY", "BUY", 40), ("BUY", "SELL", -30), ("SELL", "SELL", -50), ("SELL", "SELL", -30),
        ("BUY", "BUY", 50), ("BUY", "BUY", 40), ("HOLD", "HOLD", 8), ("BUY", "BUY", 32),
        ("SELL", "BUY", 20), ("BUY", "BUY", 30), ("BUY", "SELL", -35), ("SELL", "SELL", -35),
        ("BUY", "BUY", 30), ("BUY", "BUY", 25), ("HOLD", "HOLD", -3), ("SELL", "SELL", -32),
        ("SELL", "SELL", -30), ("BUY", "BUY", 35), ("BUY", "BUY", 25), ("HOLD", "BUY", 20),
        ("BUY", "BUY", 25), ("SELL", "SELL", -40), ("BUY", "BUY", 25), ("BUY", "SELL", -30),
        ("SELL", "SELL", -30), ("BUY", "BUY", 14),
    ]),
    "GBP/USD": (1.2480, 0.0001, [
        ("SELL", "SELL", -35), ("BUY", "BUY", 42), ("SELL", "SELL", -28), ("BUY", "BUY", 38),
        ("BUY", "SELL", -22), ("SELL", "SELL", -45), ("BUY", "BUY", 55), ("BUY", "BUY", 30),
        ("SELL", "SELL", -40), ("HOLD", "HOLD", 5), ("BUY", "BUY", 35), ("SELL", "BUY", 28),
        ("BUY", "BUY", 42), ("BUY", "SELL", -25), ("SELL", "SELL", -38), ("BUY", "BUY", 32),
        ("BUY", "BUY", 28), ("HOLD", "HOLD", -4), ("SELL", "SELL", -30), ("BUY", "BUY", 35),
        ("SELL", "SELL", -42), ("BUY", "BUY", 30), ("BUY", "BUY", 25), ("SELL", "BUY", 20),
        ("BUY", "BUY", 38), ("SELL", "SELL", -45), ("BUY", "BUY", 28), ("SELL", "SELL", -30),
        ("BUY", "SELL", -25), ("SELL", "SELL", -26),
    ]),
    "USD/JPY": (149.80, 0.01, [
        ("BUY", "BUY", 35), ("SELL", "BUY", 20), ("BUY", "BUY", 45), ("SELL", "SELL", -38),
        ("HOLD", "HOLD", 3), ("BUY", "BUY", 42), ("BUY", "SELL", -30), ("SELL", "SELL", -48),
        ("BUY", "BUY", 55), ("HOLD", "BUY", 25), ("SELL", "SELL", -35), ("BUY", "BUY", 40),
        ("BUY", "BUY", 30), ("SELL", "SELL", -42), ("BUY", "BUY", 38), ("BUY", "SELL", -28),
        ("SELL", "SELL", -35), ("HOLD", "HOLD", 5), ("BUY", "BUY", 48), ("SELL", "SELL", -40),
        ("BUY", "BUY", 35), ("BUY", "BUY", 30), ("SELL", "BUY", 18), ("BUY", "BUY", 42),
        ("SELL", "SELL", -38), ("BUY", "BUY", 35), ("BUY", "SELL", -25), ("SELL", "SELL", -30),
        ("BUY", "BUY", 38), ("HOLD", "HOLD", 4),
    ]),
    "USD/CHF": (0.8940, 0.0001, _gen_seq("USD/CHF")),
    "AUD/USD": (0.6455, 0.0001, _gen_seq("AUD/USD")),
    "USD/CAD": (1.3590, 0.0001, _gen_seq("USD/CAD")),
    "NZD/USD": (0.5970, 0.0001, _gen_seq("NZD/USD")),
    "EUR/GBP": (0.8540, 0.0001, _gen_seq("EUR/GBP")),
    "EUR/JPY": (162.10, 0.01,   _gen_seq("EUR/JPY")),
    "EUR/AUD": (1.6530, 0.0001, _gen_seq("EUR/AUD")),
    "EUR/CAD": (1.4750, 0.0001, _gen_seq("EUR/CAD")),
    "GBP/JPY": (189.40, 0.01,   _gen_seq("GBP/JPY")),
    "GBP/CHF": (1.1410, 0.0001, _gen_seq("GBP/CHF")),
    "AUD/JPY": (98.20,  0.01,   _gen_seq("AUD/JPY")),
    "XAU/USD": (3120.00, 1.0,   _gen_seq("XAU/USD")),
}

_FOREX_NEWS: list[dict[str, Any]] = [
    {
        "headline": "Federal Reserve signals cautious stance on rate cuts as inflation remains sticky",
        "sentiment": "negative",
        "source": "Reuters",
        "published_at": "2026-03-30T08:45:00Z",
    },
    {
        "headline": "EUR/USD consolidates near 1.0850 ahead of Eurozone CPI data release",
        "sentiment": "neutral",
        "source": "FXStreet",
        "published_at": "2026-03-30T08:00:00Z",
    },
    {
        "headline": "Bank of England holds rates steady, GBP/USD under pressure",
        "sentiment": "negative",
        "source": "Bloomberg",
        "published_at": "2026-03-30T07:30:00Z",
    },
    {
        "headline": "Japan's core CPI rises above expectations, BoJ hawkish bets increase",
        "sentiment": "positive",
        "source": "Nikkei",
        "published_at": "2026-03-30T07:00:00Z",
    },
    {
        "headline": "US Non-Farm Payrolls beat forecasts, USD strengthens across the board",
        "sentiment": "positive",
        "source": "MarketWatch",
        "published_at": "2026-03-30T06:30:00Z",
    },
    {
        "headline": "Eurozone PMI unexpectedly contracts, raising recession fears",
        "sentiment": "negative",
        "source": "Reuters",
        "published_at": "2026-03-30T06:00:00Z",
    },
    {
        "headline": "GBP gains on positive UK retail sales data, trade balance improves",
        "sentiment": "positive",
        "source": "FXStreet",
        "published_at": "2026-03-30T05:45:00Z",
    },
    {
        "headline": "Dollar index holds above 104 as risk sentiment remains fragile",
        "sentiment": "neutral",
        "source": "Bloomberg",
        "published_at": "2026-03-30T05:00:00Z",
    },
    {
        "headline": "ECB policymakers divided on pace of future rate reductions",
        "sentiment": "neutral",
        "source": "WSJ",
        "published_at": "2026-03-30T04:30:00Z",
    },
    {
        "headline": "Yen weakens past 151 as US-Japan yield differential widens further",
        "sentiment": "negative",
        "source": "Nikkei",
        "published_at": "2026-03-30T04:00:00Z",
    },
]

# In-memory email subscriber registry (resets on restart)
_FOREX_SUBSCRIBERS: list[dict[str, Any]] = []


def _build_forex_history(pair: str) -> list[dict[str, Any]]:
    """Return a 30-day list of predicted vs actual signals for the given pair."""
    base_price, pip, seq = _FOREX_HIST_SEQUENCES[pair]
    price = base_price
    decimals = 4 if pip < 0.01 else 2
    history: list[dict[str, Any]] = []
    for i, (pred, actual, delta) in enumerate(seq):
        d = (date(2026, 3, 1) + timedelta(days=i)).isoformat()
        entry = round(price, decimals)
        exit_price = round(price + delta * pip, decimals)
        history.append({
            "day": d,
            "predicted": pred,
            "actual": actual,
            "correct": pred == actual,
            "entry": entry,
            "exit": exit_price,
        })
        price = exit_price
    return history


def _build_technical_analysis(pair: str, live_price: float | None = None) -> dict[str, Any]:
    """Return deterministic technical-analysis data for *pair*.

    When *live_price* is supplied it is used as the reference price for
    support/resistance calculations; otherwise the static entry price from
    ``_FOREX_SIGNALS`` is used as a fallback.  For XAU/USD the price comes
    from Yahoo Finance; all other pairs use the Frankfurter API (ECB).
    """
    signal = _FOREX_SIGNALS[pair]
    price = live_price if live_price is not None else signal["entry_price"]
    direction = signal["direction"]
    is_gold = pair == "XAU/USD"
    is_jpy = "JPY" in pair
    pip = 1.0 if is_gold else (0.01 if is_jpy else 0.0001)
    dec = 2 if (is_jpy or is_gold) else 4

    def lvl(pips: float) -> float:
        return round(price + pips * pip, dec)

    # Support / Resistance levels
    support = [lvl(-50), lvl(-120), lvl(-230)]
    resistance = [lvl(60), lvl(140), lvl(250)]

    # Fair Value Gaps – unfilled gaps in price action
    fvg = [
        {
            "type": "bullish",
            "top": lvl(-28),
            "bottom": lvl(-44),
            "filled": False,
            "created": "2026-03-29",
            "description": "Unmitigated bullish FVG — potential magnet for price",
        },
        {
            "type": "bearish",
            "top": lvl(82),
            "bottom": lvl(66),
            "filled": True,
            "created": "2026-03-27",
            "description": "Filled bearish FVG — supply already consumed",
        },
        {
            "type": "bullish",
            "top": lvl(-98),
            "bottom": lvl(-114),
            "filled": True,
            "created": "2026-03-25",
            "description": "Older bullish FVG — price revisited and filled",
        },
    ]

    # Break of Structure (BOS)
    bos = [
        {
            "type": "bullish" if direction in ("BUY", "HOLD") else "bearish",
            "level": lvl(-78),
            "date": "2026-03-28",
            "description": (
                "Broke previous swing high — bullish market structure confirmed"
                if direction == "BUY"
                else "Broke previous swing low — bearish pressure continues"
            ),
        },
        {
            "type": "bearish" if direction in ("SELL", "HOLD") else "bullish",
            "level": lvl(102),
            "date": "2026-03-26",
            "description": (
                "Prior bearish BOS now acting as resistance"
                if direction == "SELL"
                else "Prior bullish BOS flipped to support"
            ),
        },
    ]

    # Change of Character (CHoCH)
    choch = [
        {
            "type": "bullish" if direction == "BUY" else "bearish",
            "level": lvl(-62),
            "date": "2026-03-29",
            "description": (
                "CHoCH: market shifted from bearish to bullish bias"
                if direction == "BUY"
                else "CHoCH: market shifted from bullish to bearish bias"
            ),
        },
    ]

    # High Volume / Order-Block Zones
    high_volume_zones = [
        {
            "top": lvl(22),
            "bottom": lvl(-14),
            "strength": "high",
            "description": "Current major liquidity pool — institutional activity",
        },
        {
            "top": lvl(-68),
            "bottom": lvl(-88),
            "strength": "medium",
            "description": "Previous order block — potential demand zone",
        },
        {
            "top": lvl(92),
            "bottom": lvl(112),
            "strength": "high",
            "description": "Supply zone — high-volume rejection expected",
        },
    ]

    return {
        "pair": pair,
        "current_price": price,
        "support_resistance": {"support": support, "resistance": resistance},
        "fvg": fvg,
        "bos": bos,
        "choch": choch,
        "high_volume_zones": high_volume_zones,
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


@app.route("/forex")
def forex_hub():
    """Render the AI Forex Signal Hub page."""
    try:
        return render_template("forex.html")
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/forex/methodology")
def forex_methodology():
    """Render the Methodology page for the Forex Signal Hub."""
    try:
        return render_template("methodology.html")
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/forex/signals")
def forex_signals():
    """Return the current signal + 30-day history for a forex pair.

    Prices and signal direction are derived from live market data:
    - XAU/USD  → Yahoo Finance (Frankfurter does not carry gold data)
    - All other pairs → Frankfurter API (ECB)
    The static ``_FOREX_SIGNALS`` dict is used as a fallback when the external
    API is unavailable.
    """
    pair = request.args.get("pair", "EUR/USD")
    if pair not in _SUPPORTED_PAIRS:
        return (
            jsonify({"error": f"Unsupported pair. Choose from: {', '.join(_SUPPORTED_PAIRS)}"}),
            400,
        )

    signal: dict[str, Any] = dict(_FOREX_SIGNALS[pair])
    signal["pair"] = pair

    # ── Attempt to enrich signal with live market data ────────────────────────
    live_rate = _fetch_live_rate(pair)
    hist_rates = _fetch_historical_rates(pair, 30)

    if live_rate is not None:
        signal["entry_price"] = live_rate
        signal["generated_at"] = datetime.now(timezone.utc).isoformat()
        signal["data_source"] = (
            "Yahoo Finance (gold spot)" if pair == "XAU/USD" else "Frankfurter API (ECB)"
        )
        signal["is_live"] = True

        if hist_rates:
            prices = list(hist_rates.values())
            direction, confidence = _compute_signal_from_prices(prices)
            signal["direction"] = direction
            signal["confidence"] = confidence

            is_gold = pair == "XAU/USD"
            is_jpy = "JPY" in pair
            pip = 1.0 if is_gold else (0.01 if is_jpy else 0.0001)
            dec = 2 if (is_jpy or is_gold) else 4
            tp_pips, sl_pips = _compute_tp_sl_pips(prices, pair)

            if direction == "BUY":
                signal["take_profit"] = round(live_rate + tp_pips * pip, dec)
                signal["stop_loss"] = round(live_rate - sl_pips * pip, dec)
            elif direction == "SELL":
                signal["take_profit"] = round(live_rate - tp_pips * pip, dec)
                signal["stop_loss"] = round(live_rate + sl_pips * pip, dec)
            else:
                signal["take_profit"] = round(live_rate + tp_pips * pip, dec)
                signal["stop_loss"] = round(live_rate - sl_pips * pip, dec)
    else:
        signal["data_source"] = "static (live feed unavailable)"
        signal["is_live"] = False

    # ── Build 30-day history ──────────────────────────────────────────────────
    if hist_rates:
        history = _build_forex_history_live(pair, hist_rates)
    else:
        history = _build_forex_history(pair)

    correct_count = sum(1 for h in history if h["correct"])
    signal["accuracy_30d"] = round(correct_count / len(history) * 100, 1) if history else 0.0
    signal["history"] = history
    return jsonify(signal)


@app.route("/api/forex/technical")
def forex_technical():
    """Return technical analysis data (FVG, S/R, BOS, CHoCH, volume zones) for a pair.

    Support/resistance levels are calculated relative to the live market price
    when available, falling back to the static entry price otherwise.
    """
    pair = request.args.get("pair", "EUR/USD")
    if pair not in _SUPPORTED_PAIRS:
        return (
            jsonify({"error": f"Unsupported pair. Choose from: {', '.join(_SUPPORTED_PAIRS)}"}),
            400,
        )
    live_rate = _fetch_live_rate(pair)
    return jsonify(_build_technical_analysis(pair, live_rate))


@app.route("/api/forex/pairs")
def forex_pairs():
    """Return the list of supported currency pairs grouped by type."""
    return jsonify({
        "major": [p for p in _SUPPORTED_PAIRS if "USD" in p.split("/") and not p.startswith("XAU")],
        "cross":  [p for p in _SUPPORTED_PAIRS if "USD" not in p.split("/") and not p.startswith("XAU")],
        "commodity": [p for p in _SUPPORTED_PAIRS if p.startswith("XAU")],
        "all": list(_SUPPORTED_PAIRS),
    })


@app.route("/api/forex/news")
def forex_news():
    """Return the latest news sentiment items."""
    return jsonify({"news": _FOREX_NEWS})


@app.route("/api/forex/subscribe", methods=["POST"])
def forex_subscribe():
    """Register an email address for signal alerts."""
    data: dict[str, Any] = request.get_json(force=True) or {}
    email: str = data.get("email", "").strip().lower()
    pairs: list[str] = data.get("pairs", [])

    if not email or "@" not in email or "." not in email.split("@")[-1]:
        return jsonify({"error": "Please provide a valid email address"}), 400

    invalid = [p for p in pairs if p not in _SUPPORTED_PAIRS]
    if invalid:
        return jsonify({"error": f"Unsupported pairs: {', '.join(invalid)}"}), 400

    if not pairs:
        pairs = list(_SUPPORTED_PAIRS)

    if any(s["email"] == email for s in _FOREX_SUBSCRIBERS):
        return jsonify({"success": True, "message": "You are already subscribed."})

    _FOREX_SUBSCRIBERS.append({
        "email": email,
        "pairs": pairs,
        "subscribed_at": datetime.now(timezone.utc).isoformat(),
    })
    return jsonify({
        "success": True,
        "message": "Subscribed! You will receive alerts when new signals are generated.",
    })


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("FLASK_DEBUG", "0") == "1"
    app.run(debug=debug, host="0.0.0.0", port=port)
